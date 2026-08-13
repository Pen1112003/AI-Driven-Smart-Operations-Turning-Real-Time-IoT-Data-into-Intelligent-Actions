import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Page setup - 0.75 inch margins for professional look
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # Style definitions
    COLOR_PRIMARY = RGBColor(21, 128, 61)    # Forest / Emerald Green (#15803D)
    COLOR_NAVY = RGBColor(15, 23, 42)        # Charcoal / Dark Slate (#0F172A)
    COLOR_ACCENT = RGBColor(2, 132, 199)     # Sky Blue (#0284C7)
    COLOR_MUTED = RGBColor(71, 85, 105)      # Slate Gray (#475569)
    COLOR_GOLD = RGBColor(180, 83, 9)        # Ministerial Gold / Amber (#B45309)
    COLOR_RED = RGBColor(220, 38, 38)        # Red Alert (#DC2626)

    # 1. HEADER & COVER BANNER
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_event = p_top.add_run("SEAL HACKATHON SUMMER 2026 • OFFICIAL TECHNICAL MANUAL\n")
    r_event.font.name = "Arial"
    r_event.font.size = Pt(10.5)
    r_event.font.bold = True
    r_event.font.color.rgb = COLOR_GOLD

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("AI-DRIVEN SMART OPERATIONS:\nFROM REAL-TIME IOT DATA TO INTELLIGENT DECISIONS")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Tài Liệu Chi Tiết: Kịch Bản Thuyết Trình 90 Phút & Nội Dung 18 Slide (Kèm Ưu/Nhược Điểm & Visual Nền Trắng Chữ Đen)")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = COLOR_MUTED

    # Metadata Table Box
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Chủ đề (Topic):", "AI-Driven Smart Operations: From Real-Time IoT Data to Intelligent Decisions"),
        ("Thời gian & Địa điểm:", "20:00 - 21:30, August 13, 2026 (90 Phút) • Google Meet (meet.google.com/zbr-tktz-bea)"),
        ("Diễn giả (Sole Speaker):", "Nguyễn Đức Tuấn (Backend Lead / System Architect, FPT Software & Google Student Ambassador)"),
        ("Kho lưu trữ mã nguồn mẫu:", "https://github.com/Pen1112003/AI-Driven-Smart-Operations-Turning-Real-Time-IoT-Data-into-Intelligent-Actions")
    ]
    for row_idx, (label, val) in enumerate(meta_data):
        c0 = meta_table.cell(row_idx, 0)
        c1 = meta_table.cell(row_idx, 1)
        c0.width = Inches(2.2)
        c1.width = Inches(4.8)
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = COLOR_NAVY

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_NAVY

    doc.add_paragraph() # Spacer

    # 2. SECTION I: TIMETABLE
    h1 = doc.add_heading("PHẦN I: TỔNG QUAN WORKSHOP & PHÂN BỔ THỜI LƯỢNG (90 PHÚT)", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY

    time_table = doc.add_table(rows=6, cols=4)
    time_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Thời Gian", "Thời Lượng", "Nội Dung Chi Tiết & Trọng Tâm Kỹ Thuật", "Diễn Giả"]
    widths = [Inches(1.2), Inches(1.0), Inches(3.8), Inches(1.0)]

    for c_idx, text in enumerate(headers):
        cell = time_table.cell(0, c_idx)
        cell.width = widths[c_idx]
        set_cell_background(cell, "15803D")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    rows_data = [
        ("20:00 - 20:10", "10 phút", "Phần 1: Khai mạc, Giới thiệu & Bối cảnh Smart Operations\n• Khởi động Hackathon & Giới thiệu lộ trình 90 phút\n• Sự chuyển dịch từ Giám sát thụ động sang AI tự hành\n• Phương pháp luận xác định Đúng/Sai (Ground Truth)", "Nguyễn Đức Tuấn"),
        ("20:10 - 20:30", "20 phút", "Phần 2: Kiến trúc Thu thập & Xử lý Luồng Dữ liệu IoT\n• Giao thức: MQTT (:1883), CoAP (:5683 UDP), Kafka/Redpanda (:9092)\n• Stream Engine: Bounded Watermarking, Tumbling (1m) & Sliding (5m) Windows", "Nguyễn Đức Tuấn"),
        ("20:30 - 20:55", "25 phút", "Phần 3: Trí tuệ Nhân tạo & AI Agents trong Vận hành Tự động\n• Chuỗi 5 Bước: Feature Extraction -> Anomaly -> RCA -> Agent -> Action\n• Isolation Forest, Autoencoder, SHAP, Causal Knowledge Graph, Guardrails", "Nguyễn Đức Tuấn"),
        ("20:55 - 21:10", "15 phút", "Phần 4: Kiến trúc Tham chiếu & Chiến Lược Hackathon\n• Blueprint 1-Click Docker Compose (FastAPI, Redis, TimescaleDB)\n• Bí quyết Pitching: Demo 'Live Anomaly Injection' 5 phút", "Nguyễn Đức Tuấn"),
        ("21:10 - 21:30", "20 phút", "Phần 5: Phiên Q&A Kỹ Thuật & Cố Vấn 1-1 Cho Thí Sinh\n• Trực tiếp tháo gỡ vướng mắc kiến trúc, streaming, AI models\n• Cảnh báo các bẫy phần cứng phổ biến và hướng dẫn bảo vệ đề tài", "Nguyễn Đức Tuấn")
    ]

    for r_idx, r_data in enumerate(rows_data, 1):
        bg = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        for c_idx, val in enumerate(r_data):
            cell = time_table.cell(r_idx, c_idx)
            cell.width = widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_NAVY
            if c_idx == 0:
                r.font.bold = True

    doc.add_paragraph()

    # 3. SECTION II: GROUND TRUTH METHODOLOGY
    h2 = doc.add_heading("PHẦN II: NGUYÊN LÝ CỐT LÕI - PHƯƠNG PHÁP XÁC ĐỊNH ĐÚNG/SAI (GROUND TRUTH)", level=1)
    h2.runs[0].font.color.rgb = COLOR_PRIMARY

    p_intro = doc.add_paragraph()
    p_intro.add_run("Một trong những thách thức kỹ thuật cốt lõi trong Smart Operations là: ").font.size = Pt(10)
    r_q = p_intro.add_run("Làm sao để hệ thống phân định chính xác giữa Nhiễu cảm biến đơn lẻ (Sensor Noise) và Sự cố vật lý thực tế (True System Anomaly) mà không gây ra báo động giả?")
    r_q.font.bold = True
    r_q.font.color.rgb = COLOR_GOLD

    doc.add_paragraph("Hệ thống áp dụng Ma trận Kiểm định Đa tầng (Multi-Layer Verification Matrix):", style='List Bullet')
    doc.add_paragraph("Tầng 1 - Giới hạn vật lý (Physical Bounds): Kiểm tra ngưỡng cực trị vật lý của cảm biến (0 <= Humidity <= 100%, 800 <= Pressure <= 1080 hPa). Bất kỳ giá trị nào ngoài dải sẽ bị loại trừ ngay.", style='List Bullet 2')
    doc.add_paragraph("Tầng 2 - Tương quan đa biến (Multivariate Correlation): Nếu chỉ 1 cảm biến nhảy vọt đơn độc trong khi các cảm biến liên đới xung quanh đứng yên -> Xác định LỖI CẢM BIẾN (Sensor Fault), không kích hoạt dừng máy.", style='List Bullet 2')
    doc.add_paragraph("Tầng 3 - Mô hình học máy (Machine Learning Deviations): Isolation Forest đánh giá Anomaly Score s(x, n) > 0.65 và Autoencoder tính sai số tái tạo MSE = ||x - x_hat||^2 > Threshold.", style='List Bullet 2')
    doc.add_paragraph("Tầng 4 - Lớp kiểm soát an toàn (Safety Guardrails): Bắt buộc mọi hành động của AI Agent phải đi qua bộ quy tắc an toàn tiền định trước khi phát lệnh điều khiển vật lý.", style='List Bullet 2')

    doc.add_paragraph()

    # 4. SECTION III: THE 18 SLIDES DETAILS
    h3 = doc.add_heading("PHẦN III: NỘI DUNG CHI TIẾT 18 SLIDE THUYẾT TRÌNH (SLIDE DECK & SPEAKER SCRIPT)", level=1)
    h3.runs[0].font.color.rgb = COLOR_PRIMARY

    SLIDES = [
        {
            "num": 1,
            "title": "TITLE SLIDE: AI-DRIVEN SMART OPERATIONS",
            "bullets": [
                "Topic: AI-Driven Smart Operations: From Real-Time IoT Data to Intelligent Decisions",
                "Speaker: Nguyen Duc Tuan (Backend Lead / System Architect, FPT Software)",
                "Event: SEAL Hackathon Summer 2026 (90 Minutes, Google Meet)",
                "Focus: Distributed Ingestion, Stream Analytics, Anomaly Detection & AI Agents"
            ],
            "details": "Chuyển dịch toàn diện từ Giám sát thụ động (Passive Telemetry) sang Vận hành thông minh tự động (AI-Driven Smart Operations). Tích hợp toàn diện chuỗi: IoT Ingestion -> Event Streaming -> Time-Series Storage -> AI Anomaly -> AI Agent Action.",
            "visual": "Màu nền: Trắng tinh khiết (#FFFFFF), Tiêu đề chữ đen xanh lá (#15803D, Outfit Bold 28pt), biểu tượng bánh răng kết hợp nút mạng IoT đối xứng trang trọng.",
            "notes": "Nhiệt liệt chào mừng các đội thi của SEAL Hackathon Summer 2026. Tối nay chúng ta sẽ cùng giải bài toán: Làm thế nào biến dữ liệu cảm biến thời gian thực thành hành động vận hành thông minh tự động!"
        },
        {
            "num": 2,
            "title": "WORKSHOP AGENDA: 90-MINUTE ROADMAP",
            "bullets": [
                "1. Overview & Evolution of Smart Operations (10 Mins)",
                "2. Real-Time IoT Streaming Pipeline Architecture (20 Mins)",
                "3. AI Decision Engine & Agentic Workflows (25 Mins)",
                "4. End-to-End Reference Architecture & Hackathon Strategy (15 Mins)",
                "5. Live Technical Q&A & Mentor Guidance (20 Mins)"
            ],
            "details": "Lộ trình 5 phân đoạn từ nền tảng giao thức đến mô hình AI và kỹ năng bảo vệ bài thi trước Ban Giám Khảo.",
            "visual": "Lưới 2 cột nền trắng (#FFFFFF), viền thẻ xanh dương (#0284C7) và xanh lá (#15803D), thời lượng hiển thị huy hiệu vàng hổ phách (#B45309).",
            "notes": "Agenda 5 phần được thiết kế mạch lạc từ tư duy kiến trúc đến triển khai code thực tế cho cuộc thi."
        },
        {
            "num": 3,
            "title": "THE EVOLUTION OF OPERATIONAL MONITORING",
            "bullets": [
                "Phase 1: Passive Monitoring ('What happened?') - Thu thập dữ liệu cảm biến & vẽ biểu đồ tĩnh. Nhược điểm: Phụ thuộc 100% vào con người quan sát 24/7.",
                "Phase 2: Reactive Alerting ('Is it broken right now?') - Cài ngưỡng tĩnh IF-THEN. Nhược điểm: Bão cảnh báo giả, không bắt được thoái hóa từ từ.",
                "Phase 3: AI-Driven Smart Operations ('Why & How to fix?') - Tự động phát hiện sớm, chẩn đoán nguyên nhân gốc rễ (RCA) và tự động thực thi hành động."
            ],
            "details": "Khái niệm cốt lõi: Phân biệt 3 giai đoạn tiến hóa của hệ thống giám sát công nghiệp và lợi thế cạnh tranh của AI Agent.",
            "visual": "Bảng so sánh 3 cột nền trắng (Cột 1 xám nhạt #F1F5F9, Cột 2 vàng cam nhạt #FEF3C7, Cột 3 xanh lá nhạt #DCFCE7), chữ đen than chì đậm nét.",
            "notes": "Nếu sản phẩm Hackathon của bạn chỉ hiển thị biểu đồ, bạn đang ở Phase 1. Để thắng giải, bạn cần đưa sản phẩm sang Phase 3: AI chủ động ra quyết định."
        },
        {
            "num": 4,
            "title": "REAL-TIME IOT DATA PIPELINE ARCHITECTURE",
            "bullets": [
                "Layer 1 - Edge & Ingestion: Sensor Nodes -> MQTT (:1883) & CoAP (:5683 UDP)",
                "Layer 2 - Distributed Buffer: Redpanda / Apache Kafka (:9092) chống nghẽn và đảm bảo Zero Data Loss",
                "Layer 3 - Stream Engine: Python Async / Flink xử lý Watermarking & Windowing",
                "Layer 4 - Dual Storage & AI: Redis Hot Cache (<1ms) + TimescaleDB Hypertables nén 90%"
            ],
            "details": "Kiến trúc luồng dữ liệu 4 tầng phân tách rõ ràng giữa tốc độ thu thập và tốc độ suy luận của mô hình trí tuệ nhân tạo.",
            "visual": "Sơ đồ khối 4 tầng từ trên xuống nền trắng (#FFFFFF), viền hộp than chì (#334155), mũi tên luồng dữ liệu xanh dương (#0284C7), badge giao thức xanh lá (#15803D).",
            "notes": "Dữ liệu IoT có tính chất High-velocity. Bạn bắt buộc phải có lớp Buffer như Kafka hay Redis Streams để chống nghẽn hệ thống."
        },
        {
            "num": 5,
            "title": "INGESTION TECHNOLOGIES: MQTT, COAP & KAFKA/REDPANDA",
            "bullets": [
                "MQTT (1883): Pub/Sub trên TCP, header 2B siêu nhẹ, QoS 0/1/2, Keepalive, LWT. Lưu ý: Dùng QoS 1 + event_time.",
                "CoAP (5683 UDP): RESTful nhị phân chạy UDP, header 4B, tối ưu pin. Lưu ý: Phù hợp cho node cảm biến pin yếu.",
                "Redpanda vs Kafka (9092): Commit log phân tán bất biến. Redpanda viết bằng C++, zero JVM, p99 latency thấp 10x, ngốn <200MB RAM. Lưu ý: Dùng station_id làm Partition Key!"
            ],
            "details": "So sánh chi tiết 3 công nghệ Ingestion chủ lực: Định nghĩa, Ưu điểm, Nhược điểm và Lưu ý cấu hình tránh mất gói dữ liệu.",
            "visual": "Lưới 3 cột nền trắng, thẻ bo góc viền màu phân biệt (Xanh lá, Xanh dương, Vàng cam), bảng mini 2 hàng (Xanh: Ưu điểm | Đỏ: Lưu ý Gotchas).",
            "notes": "MQTT giúp cảm biến gửi tin siêu nhẹ, CoAP tối ưu pin, còn Redpanda C++ là cứu cánh giúp bạn chạy Kafka mượt mà trên laptop."
        },
        {
            "num": 6,
            "title": "STREAM PROCESSING: WATERMARKING & TIME WINDOWING",
            "bullets": [
                "Event-Time vs Processing-Time: Luôn xử lý theo giờ tạo dữ liệu của cảm biến (Event-Time) để đảm bảo tính đúng đắn.",
                "Bounded Watermarking: Watermark(t) = max(Event_Time) - T_delay. Chấp nhận chờ gói tin trễ 5s. Late data đưa vào Side-Output.",
                "Tumbling Window (1m): Cố định không chồng lấn cho thống kê chu kỳ cơ sở.",
                "Sliding Window (5m, trượt 1m): Cửa sổ trượt liên tục tính đạo hàm biến thiên khí áp Delta P / Delta t báo bão."
            ],
            "details": "Giải thuật xử lý dữ liệu trễ mạng và mô hình tính toán cửa sổ thời gian thực phục vụ phát hiện bất thường.",
            "visual": "Nền trắng (#FFFFFF), biểu đồ trục thời gian minh họa Watermark trễ và 2 sơ đồ cửa sổ Tumbling (nối đuôi) vs Sliding (gối đầu 80%).",
            "notes": "Watermarking là vũ khí giải quyết triệt để bài toán mạng chập chờn, giúp cửa sổ phân tích thời gian thực không bao giờ bị tính sai."
        },
        {
            "num": 7,
            "title": "TIME-SERIES STORAGE: REDIS HOT CACHE VS TIMESCALEDB",
            "bullets": [
                "Redis Streams (Hot Cache < 1ms): Bộ nhớ RAM lưu trữ trạng thái tức thời phục vụ Feature Vectors cho AI Model suy luận. Lưu ý: Cấu hình XADD MAXLEN ~ 1000 chống tràn RAM.",
                "TimescaleDB (Cold Time-Series Engine): Mở rộng PostgreSQL thành CSDL chuỗi thời gian qua Hypertables. Tự động nén Columnar Compression 90%+ dung lượng đĩa và hỗ trợ 100% SQL chuẩn.",
                "InfluxDB: CSDL chuyên dụng TSM, ghi cực nhanh nhưng hạn chế High Cardinality và phải học ngôn ngữ Flux mới."
            ],
            "details": "Chiến lược phân tầng lưu trữ Hot/Cold tối ưu hóa chi phí phần cứng và tốc độ truy vấn của hệ thống.",
            "visual": "Bảng so sánh phân tầng Hot (Vàng cam nhạt #FFFBEB) ở trên và Cold (Xanh biển nhạt #F0F9FF) ở dưới, chữ đen rõ nét, cột so sánh Latency, Retention, Compression.",
            "notes": "Mô hình Hot/Cold phân tách rõ ràng: Redis phục vụ AI suy luận dưới 5ms, còn TimescaleDB lưu trữ toàn bộ lịch sử phục vụ đào tạo lại mô hình."
        },
        {
            "num": 8,
            "title": "THE 5-STEP AI DECISION PIPELINE OVERVIEW",
            "bullets": [
                "1. Feature Extraction: Rolling statistics (Mean, StdDev) & FFT frequency harmonics.",
                "2. Anomaly Detection: Isolation Forest & Deep Autoencoder Reconstruction Error (MSE > Threshold).",
                "3. Root Cause Analysis (RCA): SHAP feature attribution & Causal Knowledge Graphs chỉ rõ linh kiện lỗi.",
                "4. AI Agent (LLM + Guardrails): Lập luận logic có bọc Deterministic Safety Guardrails chống ảo giác.",
                "5. Automated Action Execution: Closed-loop MQTT control dispatches & bảng điều hành khuyến nghị tự động."
            ],
            "details": "Chuỗi 5 bước khép kín chuyển đổi hoàn toàn tín hiệu cảm biến thô thành hành động điều hành tự động.",
            "visual": "Quy trình 5 khối hình chữ nhật liên hoàn dạng Chevron nối bằng mũi tên xanh lá (#15803D), số thứ tự in đậm, chữ đen than chì.",
            "notes": "Đây là linh hồn của buổi workshop. 5 bước nối tiếp chuyển hóa hoàn toàn dữ liệu cảm biến thô thành hành động tự động."
        },
        {
            "num": 9,
            "title": "STEP 1 & 2: FEATURE EXTRACTION & ANOMALY DETECTION",
            "bullets": [
                "Feature Extraction: Miền thời gian (Rolling Mean/Variance, Z-Score) + Miền tần số (FFT trích xuất Harmonic Peaks rung động động cơ).",
                "Isolation Forest (Scikit-Learn): Phân lập điểm dị biệt qua cây ngẫu nhiên O(n log n). Cực nhanh, chạy trên CPU, không cần nhãn. Đặt contamination=0.03.",
                "Autoencoder Neural Network (PyTorch): Nén Latent Space rồi tái tạo. Bất thường khi Reconstruction Error MSE = ||x - x_hat||^2 > Threshold. Export sang ONNX Runtime để suy luận < 5ms."
            ],
            "details": "So sánh toán học và giải thuật giữa Machine Learning cổ điển (Isolation Forest) và Deep Learning (Autoencoder).",
            "visual": "Chia đôi màn hình (Grid-2) nền trắng, bên trái là biểu đồ cây phân lập Isolation Forest, bên phải là mô hình Autoencoder thắt nút cổ chai (Hourglass).",
            "notes": "Isolation Forest rất nhẹ và sẵn sàng dùng ngay. Nếu dùng Deep Learning Autoencoder, hãy export ra ONNX Runtime để suy luận nhanh."
        },
        {
            "num": 10,
            "title": "STEP 3: ROOT CAUSE ANALYSIS (RCA) & EXPLAINABILITY",
            "bullets": [
                "SHAP (Shapley Additive exPlanations): Đo lường mức đóng góp % của từng cảm biến vào điểm bất thường (Ví dụ: Rung trục +78%, Áp suất +14%, Nhiệt độ +8%). Giúp giải thích tường minh (XAI).",
                "Causal Knowledge Graphs: Đồ thị DAG mô tả quan hệ nhân quả vật lý (Tắc van A -> Tăng áp B -> Quá nhiệt C). Phân biệt chính xác 'Nguyên nhân gốc' vs 'Triệu chứng lan truyền'."
            ],
            "details": "Phương pháp chẩn đoán nguyên nhân gốc rễ và ứng dụng Explainable AI trong bảo trì dự đoán.",
            "visual": "Nền trắng (#FFFFFF), thẻ trái có biểu đồ thanh ngang SHAP Waterfall, thẻ phải có đồ thị nút mạng (Nodes) nối mũi tên đỏ chỉ vào nút gốc.",
            "notes": "Đừng chỉ báo lỗi chung chung! Hãy dùng SHAP để chỉ rõ cho Giám khảo thấy linh kiện nào đang gây ra 80% nguyên nhân sự cố."
        },
        {
            "num": 11,
            "title": "STEP 4: AI AGENTS, LLMS & SAFETY GUARDRAILS",
            "bullets": [
                "LLM Reasoning Engine (Gemini / OpenAI): Tiếp nhận chỉ số cảm biến + kết quả RCA SHAP + cẩm nang SOPs để lập luận và chọn tool (Function Calling).",
                "Deterministic Safety Guardrails: Chốt chặn kiểm duyệt an toàn bắt buộc trước khi phát lệnh vật lý (Ví dụ: 'Không đóng van chính nếu áp lực > 100 bar'). Có quyền PHỦ QUYẾT tuyệt đối mọi lệnh vi phạm của LLM."
            ],
            "details": "Cơ chế lập luận ngữ cảnh của AI Agent và lớp kiểm soát an toàn vật lý chống ảo giác (Hallucination).",
            "visual": "Sơ đồ 3 lớp bảo vệ đồng tâm nền trắng, lớp lõi LLM (Xanh dương), lớp bao bọc Guardrails (Vàng viền đỏ #DC2626), lớp ngoài Actuators (Xanh lá).",
            "notes": "Mô hình ngôn ngữ lớn rất giỏi suy luận nhưng không thể để nó tự do gửi lệnh vật lý. Lớp Safety Guardrails là chốt chặn an toàn bắt buộc."
        },
        {
            "num": 12,
            "title": "STEP 5: AUTOMATED ACTION EXECUTION & CLOSED-LOOP CONTROL",
            "bullets": [
                "Automated Closed-Loop: Gửi lệnh MQTT/API ngắt thiết bị khẩn cấp, mở trạm bơm tiêu úng tối đa công suất khi mưa > 40 mm/h, bật phun sương làm mát khi Heat Index >= 40°C.",
                "Human-in-the-Loop: Đẩy báo cáo Markdown về Telegram/Dashboard với nút bấm 1-chạm 'Xác Nhận Thực Hiện' hoặc 'Từ Chối' cho kỹ sư trưởng."
            ],
            "details": "Hai chế độ thực thi hành động tự động: Tự hành hoàn toàn cho sự cố khẩn cấp và Bán tự động có con người giám sát.",
            "visual": "Nền trắng (#FFFFFF), sơ đồ mũi tên vòng lặp khép kín tuần hoàn lớn màu xanh lá (#15803D) từ Cảm biến -> AI Agent -> Actuator -> Trạng thái an toàn.",
            "notes": "Vòng lặp hoàn tất: Dữ liệu cảm biến vào -> AI chẩn đoán -> Lệnh tự động kích hoạt thiết bị xử lý mà không cần con người can thiệp thủ công."
        },
        {
            "num": 13,
            "title": "END-TO-END HACKATHON REFERENCE ARCHITECTURE",
            "bullets": [
                "Data Producer: weather-simulator (Python asyncio phát đa trạm thời gian thực).",
                "Ingestion Gateways: weather-mqtt-bridge (:1883) & weather-coap-gateway (:5683 UDP).",
                "Event Broker: weather-redpanda (:9092) + weather-redpanda-console (:8088 Web UI).",
                "Stream & AI Engine: weather-stream-engine (Watermarking, Windowing, ONNX Inference).",
                "Backend & UI: weather-web-dashboard (FastAPI, WebSockets, Leaflet Map :8000)."
            ],
            "details": "Bản thiết kế hệ thống hoàn chỉnh 100% container hóa bằng Docker Compose cho các đội thi triển khai ngay.",
            "visual": "Sơ đồ Blueprint toàn diện nền trắng, các container đóng khung viền nét đứt (#94A3B8), cổng kết nối đóng khung vàng (#FEF3C7).",
            "notes": "Các bạn có thể chụp lại Slide này! Đây chính là Blueprint hoàn chỉnh có thể dựng ngay bằng Docker Compose trong 2 giờ."
        },
        {
            "num": 14,
            "title": "RECOMMENDED HACKATHON TECH STACK DEEP-DIVE",
            "bullets": [
                "FastAPI & Paho-MQTT: Backend phi đồng bộ siêu tốc độ xử lý hàng chục nghìn req/s, tích hợp WebSockets và tự sinh tài liệu OpenAPI.",
                "Leaflet.js & Chart.js: Trực quan hóa bản đồ địa lý và biểu đồ dòng dữ liệu trượt 60fps mà không cần framework cồng kềnh.",
                "PyTorch -> ONNX Runtime: Huấn luyện mô hình dễ dàng trong PyTorch rồi xuất file .onnx để suy luận tốc độ C++ < 5ms.",
                "Docker Compose: Đóng gói toàn bộ 8 services trong 1 file cấu hình, chạy 1-Click docker compose up -d."
            ],
            "details": "Phân tích lý do lựa chọn bộ công cụ nhẹ, hiệu năng cao, ít tốn RAM và dễ debug trong 48 giờ thi.",
            "visual": "Lưới 4 thẻ công nghệ (Grid-4) nền trắng, icon nhận diện thương hiệu, danh sách 3 tính năng cốt lõi và thanh đánh giá hiệu năng.",
            "notes": "Toàn bộ stack được lựa chọn theo tiêu chí: Nhẹ, dễ debug, không tốn RAM và khởi động trong vài giây."
        },
        {
            "num": 15,
            "title": "SECRET WEAPON: THE 'LIVE ANOMALY INJECTION' PITCH",
            "bullets": [
                "Tạo nút 'Tiêm Sự Cố' (Inject Anomaly) ngay trên giao diện Web.",
                "Kịch bản Pitching 5 phút: 1. Hệ thống bình thường (Xanh) -> 2. Bấm Tiêm Lỗi bão nhiệt đới -> 3. Đồ thị đổi màu cam, phát hiện sụt áp -> 4. AI chẩn đoán nguyên nhân trong <100ms -> 5. Tự động phát lệnh mở bơm tiêu úng thành công."
            ],
            "details": "Chiến lược thuyết trình 5 phút ấn tượng nhất: Biểu diễn sự cố xảy ra trực tiếp và AI tự động giải quyết trước mắt Ban Giám Khảo.",
            "visual": "Kịch bản phân cảnh 5 bước (5-Step Storyboard) nền trắng, nút bấm tiêm lỗi nổi bật màu đỏ rực (#DC2626), mũi tên tiến trình rõ ràng.",
            "notes": "Khi chấm thi, Giám khảo chỉ có 5 phút. Hãy cho Giám khảo thấy sự cố xảy ra và hệ thống AI tự động xử lý ngay trước mắt họ!"
        },
        {
            "num": 16,
            "title": "COMMON HACKATHON PITFALLS TO AVOID",
            "bullets": [
                "❌ Bẫy 1 - Quá sa đà vào hàn mạch phần cứng: Dành 80% thời gian ngồi hàn ESP32/Arduino -> Khắc phục: Dùng Python Simulator phát qua MQTT.",
                "❌ Bẫy 2 - Dùng mô hình LLM quá nặng: Dùng model 70B khiến API đơ 30s khi demo -> Khắc phục: Dùng Gemini-Flash + Heuristic Fallback offline.",
                "❌ Bẫy 3 - Làm Dashboard tĩnh: Chỉ vẽ biểu đồ không có hành động -> Khắc phục: Xây dựng thẻ hành động điều hành tự động qua MQTT."
            ],
            "details": "3 sai lầm kinh điển khiến các đội thi mất điểm và giải pháp khắc phục thực chiến.",
            "visual": "Bảng 3 cột so sánh lỗi sai (Đỏ nhạt #FEE2E2, Vàng cam nhạt #FEF3C7) và giải pháp chuẩn (Xanh lá nhạt #DCFCE7), chữ đen rõ ràng.",
            "notes": "Hãy nhớ: Cuộc thi đánh giá giải pháp AI-Driven Smart Operations, không phải thi hàn mạch phần cứng!"
        },
        {
            "num": 17,
            "title": "LIVE TECHNICAL Q&A & MENTORSHIP",
            "bullets": [
                "Giải đáp trực tiếp mọi thắc mắc về Streaming Pipeline, Watermarking và AI Models.",
                "Tư vấn kiến trúc 1-1 cho ý tưởng cụ thể của từng đội thi.",
                "Hướng dẫn tối ưu hóa kịch bản pitching trước Ban Giám Khảo."
            ],
            "details": "Phiên cố vấn kỹ thuật và tháo gỡ khó khăn trực tiếp giữa Diễn giả và các đội thi.",
            "visual": "Nền trắng (#FFFFFF), chữ đen than chì, biểu tượng micro và bong bóng trò chuyện (Chat Bubbles), khung thông tin liên hệ trang trọng.",
            "notes": "Bây giờ là thời gian cho các bạn! Hãy đặt bất kỳ câu hỏi nào về khó khăn kỹ thuật mà đội thi của bạn đang gặp phải."
        },
        {
            "num": 18,
            "title": "CLOSING & BEST WISHES",
            "bullets": [
                "Speaker: Nguyen Duc Tuan (Backend Lead & System Architect)",
                "Event: SEAL Hackathon Summer 2026 • Google Meet",
                "Repository: https://github.com/Pen1112003/AI-Driven-Smart-Operations-Turning-Real-Time-IoT-Data-into-Intelligent-Actions",
                "Thông điệp: Build Smart, Scale Fast, and Drive Intelligent Action!"
            ],
            "details": "Lời chúc bế mạc và công bố kho lưu trữ mã nguồn mẫu tham khảo.",
            "visual": "Slide bế mạc nền trắng trang nhã, thông điệp truyền cảm hứng in đậm màu xanh lá (#15803D), mã link repository nổi bật ở trung tâm.",
            "notes": "Chúc tất cả các đội thi SEAL Hackathon Summer 2026 xây dựng được những giải pháp Vận hành Thông minh xuất sắc!"
        }
    ]

    for item in SLIDES:
        s_card = doc.add_table(rows=5, cols=1)
        s_card.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Row 0: Header cell (Green background, white text)
        c_head = s_card.cell(0, 0)
        c_head.width = Inches(7.0)
        set_cell_background(c_head, "15803D")
        set_cell_margins(c_head, 70, 70, 100, 100)
        p_h = c_head.paragraphs[0]
        r_h = p_h.add_run(f"SLIDE {item['num']:02d}: {item['title']}")
        r_h.font.bold = True
        r_h.font.size = Pt(11)
        r_h.font.color.rgb = RGBColor(255, 255, 255)

        # Row 1: Content cell (White background, black text)
        c_body = s_card.cell(1, 0)
        c_body.width = Inches(7.0)
        set_cell_background(c_body, "FFFFFF")
        set_cell_margins(c_body, 70, 70, 120, 120)
        p_b_title = c_body.paragraphs[0]
        r_bt = p_b_title.add_run("Nội dung hiển thị trên Slide (Slide Content):")
        r_bt.font.bold = True
        r_bt.font.size = Pt(9.5)
        r_bt.font.color.rgb = COLOR_NAVY

        for b in item["bullets"]:
            p_bullet = c_body.add_paragraph(style='List Bullet')
            r_b = p_bullet.add_run(b)
            r_b.font.size = Pt(9)
            r_b.font.color.rgb = COLOR_NAVY

        # Row 2: Deep-Dive Details cell (Light slate background)
        c_details = s_card.cell(2, 0)
        c_details.width = Inches(7.0)
        set_cell_background(c_details, "F8FAFC")
        set_cell_margins(c_details, 60, 60, 120, 120)
        p_d = c_details.paragraphs[0]
        r_dt = p_d.add_run("🔍 Phân tích Khái niệm & Cơ chế Kỹ thuật: ")
        r_dt.font.bold = True
        r_dt.font.size = Pt(9)
        r_dt.font.color.rgb = COLOR_ACCENT
        r_dv = p_d.add_run(item["details"])
        r_dv.font.size = Pt(9)
        r_dv.font.color.rgb = COLOR_NAVY

        # Row 3: Visual Design cell (White background with dashed border concept)
        c_vis = s_card.cell(3, 0)
        c_vis.width = Inches(7.0)
        set_cell_background(c_vis, "F1F5F9")
        set_cell_margins(c_vis, 60, 60, 120, 120)
        p_v = c_vis.paragraphs[0]
        r_vt = p_v.add_run("🎨 Visual Design (Nền Trắng Chữ Đen & Bố Cục): ")
        r_vt.font.bold = True
        r_vt.font.size = Pt(9)
        r_vt.font.color.rgb = COLOR_PRIMARY
        r_vv = p_v.add_run(item["visual"])
        r_vv.font.size = Pt(8.5)
        r_vv.font.color.rgb = COLOR_NAVY

        # Row 4: Speaker Notes cell (Amber yellow background)
        c_notes = s_card.cell(4, 0)
        c_notes.width = Inches(7.0)
        set_cell_background(c_notes, "FEF3C7")
        set_cell_margins(c_notes, 60, 60, 120, 120)
        p_n = c_notes.paragraphs[0]
        r_nt = p_n.add_run("🎤 Lời thoại Diễn giả (Speaker Script / Notes): ")
        r_nt.font.bold = True
        r_nt.font.size = Pt(9)
        r_nt.font.color.rgb = COLOR_GOLD
        r_nv = p_n.add_run(f'"{item["notes"]}"')
        r_nv.font.italic = True
        r_nv.font.size = Pt(9)
        r_nv.font.color.rgb = COLOR_NAVY

        doc.add_paragraph() # Spacer

    # 5. SECTION IV: TECH STACK MATRIX
    h4 = doc.add_heading("PHẦN IV: BẢNG MA TRẬN SO SÁNH CÔNG NGHỆ CHUYÊN SÂU (TECH STACK MATRIX)", level=1)
    h4.runs[0].font.color.rgb = COLOR_PRIMARY

    tech_table = doc.add_table(rows=14, cols=5)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_headers = ["Tầng Kiến Trúc", "Công Nghệ", "Định Nghĩa & Cơ Chế", "Ưu Điểm & Nhược Điểm", "Lưu Ý Thực Chiến (Gotchas)"]
    t_widths = [Inches(1.2), Inches(1.1), Inches(1.8), Inches(1.6), Inches(1.3)]

    for c_idx, text in enumerate(t_headers):
        cell = tech_table.cell(0, c_idx)
        cell.width = t_widths[c_idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(8.5)

    tech_rows = [
        ("Ingestion", "MQTT (:1883)", "Giao thức Pub/Sub TCP, header 2B.", "Ưu: Tiết kiệm băng thông, Keepalive, QoS 0/1/2.\nNhược: Chạy TCP tốn pin hơn UDP.", "Dùng QoS 1 + event_time; cấu hình LWT phát hiện mất mạng."),
        ("Ingestion", "CoAP (:5683 UDP)", "RESTful nhị phân chạy trên UDP.", "Ưu: Header 4B, siêu tiết kiệm pin.\nNhược: Dễ rớt gói trên mạng chập chờn.", "Tối ưu cho node cảm biến nông nghiệp/môi trường pin yếu."),
        ("Streaming Buffer", "Redpanda (:9092)", "Commit Log phân tán viết bằng C++.", "Ưu: Zero JVM, p99 latency thấp 10x, ngốn <200MB RAM.\nNhược: Mới hơn Kafka.", "⭐ Khuyên dùng số 1 cho Hackathon; đặt key theo station_id."),
        ("Streaming Buffer", "Apache Kafka", "Distributed Event Streaming trên JVM.", "Ưu: Chuẩn công nghiệp số 1, tải cực lớn.\nNhược: Ngốn 2GB+ RAM, cấu hình nặng.", "Cần máy cấu hình mạnh; cẩn thận lỗi JVM OOM Heap."),
        ("Streaming Buffer", "Redis Streams", "Append-only Log trực tiếp trong RAM.", "Ưu: Độ trễ <1ms, cực nhẹ.\nNhược: Dữ liệu giới hạn trong RAM.", "Luôn dùng XADD MAXLEN ~ 1000 để chống tràn bộ nhớ."),
        ("Stream Engine", "Async Window Engine", "Xử lý dòng dữ liệu phi đồng bộ Python.", "Ưu: Tích hợp trực tiếp AI, kiểm soát Watermark linh hoạt.\nNhược: Tự quản lý state.", "Cấu hình allowed_delay 5s để gom dữ liệu trễ mạng."),
        ("Stream Engine", "Apache Flink", "Stateful stream engine chuẩn mực.", "Ưu: Xử lý Exactly-once, đồ thị stream lớn.\nNhược: Khởi động nặng, tốn thời gian học.", "Dùng cho hệ thống lớn cấp Enterprise."),
        ("Time-Series DB", "TimescaleDB", "PostgreSQL Extension với Hypertables.", "Ưu: SQL chuẩn, tự phân vùng, nén 90%.\nNhược: Tốn tài nguyên như Postgres.", "Cấu hình chunk 1 ngày và bật Columnar Compression."),
        ("Time-Series DB", "InfluxDB", "CSDL Time-series chuyên dụng TSM.", "Ưu: Tốc độ ghi cực nhanh.\nNhược: Lỗi High Cardinality, học ngôn ngữ Flux.", "Chỉ dùng khi đã quen thuộc với hệ sinh thái Influx."),
        ("AI Anomaly", "Isolation Forest", "Phân lập điểm dị biệt qua cây ngẫu nhiên.", "Ưu: Nhanh O(n log n), chạy CPU, Unsupervised.\nNhược: Khó học tương quan phi tuyến sâu.", "Đặt contamination=0.03, kết hợp trích xuất đặc trưng FFT."),
        ("AI Anomaly", "Autoencoder", "Mạng nơ-ron nén và tái tạo tín hiệu.", "Ưu: Học được quan hệ phức tạp, bắt lỗi mới lạ.\nNhược: Cần chuẩn hóa dữ liệu cẩn thận.", "Export sang ONNX Runtime để suy luận < 5ms."),
        ("RCA & XAI", "SHAP", "Phân bổ Shapley từ lý thuyết trò chơi.", "Ưu: Chỉ rõ % đóng góp của từng cảm biến vào lỗi.\nNhược: Tính toán lâu nếu nhiều biến.", "Dùng TreeSHAP để tính toán trong mili-giây."),
        ("AI Agent", "LLM Reasoning", "Mô hình ngôn ngữ lớn (Gemini/OpenAI).", "Ưu: Tổng hợp ngữ cảnh, phân tích đa chiều.\nNhược: Có thể bị ảo giác (Hallucination).", "Bắt buộc bọc qua Safety Guardrails trước khi phát lệnh!")
    ]

    for r_idx, r_data in enumerate(tech_rows, 1):
        bg = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        for c_idx, val in enumerate(r_data):
            cell = tech_table.cell(r_idx, c_idx)
            cell.width = t_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 50, 50, 60, 60)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_NAVY
            if c_idx == 1:
                r.font.bold = True

    doc.add_paragraph()

    # 6. SECTION V: SCORING RUBRIC
    h5 = doc.add_heading("PHẦN V: BẢNG TIÊU CHÍ CHẤM ĐIỂM & CHECKLIST DÀNH CHO THÍ SINH", level=1)
    h5.runs[0].font.color.rgb = COLOR_PRIMARY

    rubric_table = doc.add_table(rows=6, cols=3)
    rubric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_headers = ["Hạng Mục Đánh Giá", "Trọng Số", "Yêu Cầu Kỹ Thuật Đạt Chuẩn"]
    r_widths = [Inches(2.2), Inches(1.0), Inches(3.8)]

    for c_idx, text in enumerate(r_headers):
        cell = rubric_table.cell(0, c_idx)
        cell.width = r_widths[c_idx]
        set_cell_background(cell, "15803D")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    rubric_rows = [
        ("1. Kiến trúc phân tầng & Giao thức IoT", "20%", "Hỗ trợ MQTT/CoAP; có Buffer Kafka/Redpanda chống nghẽn; không thất thoát dữ liệu."),
        ("2. Stream Processing & Watermarking", "25%", "Xử lý Event-Time, Watermark trễ mạng; tính toán đúng cửa sổ Tumbling & Sliding."),
        ("3. Mô hình AI & Chẩn đoán RCA", "25%", "Trích xuất đặc trưng đa biến; phát hiện bất thường Isolation Forest/Autoencoder; chỉ ra linh kiện lỗi qua SHAP/Causal Graph."),
        ("4. Tự động hóa điều hành khép kín", "15%", "AI Agent đưa ra phương án khắc phục và phát lệnh điều khiển thiết bị tự động qua MQTT/Webhook có kiểm soát an toàn."),
        ("5. Giao diện trực quan & Kỹ năng Demo", "15%", "Dashboard trực quan hóa thời gian thực; trình diễn kịch bản Tiêm Lỗi (Live Anomaly Injection) thuyết phục trước Giám khảo.")
    ]

    for r_idx, r_data in enumerate(rubric_rows, 1):
        bg = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        for c_idx, val in enumerate(r_data):
            cell = rubric_table.cell(r_idx, c_idx)
            cell.width = r_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_NAVY
            if c_idx == 1:
                r.font.bold = True

    output_path = "/Users/penpen1112003/AI-Driven_Smart_Operations- _From_Real-Time_IoT_Data_to_Intelligent_Decisions/WORKSHOP_AI_DRIVEN_SMART_OPERATIONS_SLIDE_DECK.docx"
    doc.save(output_path)
    print(f"Successfully generated expanded Word document at: {output_path}")

if __name__ == "__main__":
    create_document()
